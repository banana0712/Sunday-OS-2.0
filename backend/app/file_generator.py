"""
SundayOS 文件生成模块
- Word 报告（学术/简洁两种风格）
- 数据图表（matplotlib）
- 免费配图（LoremFlickr — 稳定可靠）
"""
import io
import base64
import hashlib
import logging
import httpx
from datetime import datetime
from zoneinfo import ZoneInfo

logger = logging.getLogger(__name__)
TZ = ZoneInfo("Asia/Shanghai")


# ============================================================
# 免费配图 — LoremFlickr（稳定、免费、无 API key）
# ============================================================

def get_image_url(keyword: str, width: int = 800, height: int = 400) -> str:
    """
    获取免费配图 URL。
    使用 LoremFlickr — 最稳定的免费图源，支持关键词搜索。
    同一关键词返回同一张图（通过缓存 seed）。
    """
    keyword_clean = keyword.lower().strip().replace(" ", ",")[:50]
    # 使用 seed hash 确保同一关键词总返回同一张图
    seed = int(hashlib.md5(keyword.encode()).hexdigest()[:8], 16) % 10000
    return f"https://loremflickr.com/{width}/{height}/{keyword_clean}?random={seed}"


def get_unsplash_url(keyword: str, width: int = 800, height: int = 400) -> str:
    """兼容旧接口"""
    return get_image_url(keyword, width, height)


# ============================================================
# Word 报告生成
# ============================================================

async def generate_word_report(
    llm_client,
    topic: str,
    user_context: str = "",
    style: str = "auto",  # auto / academic / brief
) -> tuple[bytes, str]:
    """
    生成 Word 报告。
    返回 (docx_bytes, filename)
    
    流程：搜索资料 → 生成大纲 → 逐章节撰写 → 格式化 Word
    """
    from app.config import settings as app_settings
    from app.search import search_web, format_search_results

    model = app_settings.llm_model

    # ── 1. 搜索资料 ──
    search_results = search_web(topic, max_results=5)
    search_text = format_search_results(search_results)

    # ── 2. 生成大纲 ──
    style_guide = ""
    if style == "academic":
        style_guide = "学术正式风格，包含摘要、引言、方法、讨论、结论、参考文献"
    elif style == "brief":
        style_guide = "简洁实用风格，要点列表 + 小标题 + 简短段落"
    else:
        style_guide = "根据主题自动判断：学术类用正式风格，日常类用简洁风格"

    outline_prompt = f"""请为以下主题生成一份报告大纲。

主题：{topic}
用户背景：{user_context or "普通读者"}
风格要求：{style_guide}

参考资料：
{search_text[:1500]}

请输出一个 JSON 格式的大纲：
{{
  "title": "报告标题",
  "style": "academic 或 brief",
  "sections": [
    {{"heading": "章节标题", "key_points": ["要点1", "要点2"]}}
  ]
}}

只输出 JSON。"""

    resp = await llm_client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": outline_prompt}],
        temperature=0.7, max_tokens=800,
    )
    import json, re
    text = resp.choices[0].message.content.strip()
    json_match = re.search(r'\{[\s\S]*\}', text)
    if json_match:
        outline = json.loads(json_match.group())
    else:
        outline = {"title": topic, "style": "brief", "sections": [{"heading": topic, "key_points": ["概述"]}]}

    title = outline.get("title", topic)
    report_style = outline.get("style", "brief")
    sections = outline.get("sections", [])

    # ── 3. 逐章节撰写 ──
    full_content = []
    for i, sec in enumerate(sections):
        heading = sec.get("heading", "")
        key_points = sec.get("key_points", [])
        kp_text = "\n".join([f"- {p}" for p in key_points])

        section_prompt = f"""请撰写报告章节。

报告标题：{title}
当前章节：{heading}
章节要点：
{kp_text}

参考资料：{search_text[:1000]}
用户背景：{user_context or "普通读者"}
风格：{'学术正式' if report_style == 'academic' else '简洁实用'}

请直接写正文，2-4段，不要写"第X章"之类的编号。语言自然流畅。"""

        resp = await llm_client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": section_prompt}],
            temperature=0.7, max_tokens=600,
        )
        body = resp.choices[0].message.content.strip()
        full_content.append({"heading": heading, "body": body})

    # ── 4. 生成参考文献（学术风格）──
    references = []
    if report_style == "academic":
        ref_prompt = f"""根据以下搜索资料，提取3-5条参考文献，格式为：作者. 标题. 来源, 年份.

{search_text[:1000]}

每行一条，不要编号。"""
        resp = await llm_client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": ref_prompt}],
            temperature=0.5, max_tokens=300,
        )
        refs_text = resp.choices[0].message.content.strip()
        references = [r.strip() for r in refs_text.split("\n") if r.strip()]

    # ── 5. 构建 Word 文档 ──
    docx_bytes = _build_docx(title, full_content, references, report_style, search_results)

    # 文件名
    safe_title = re.sub(r'[\\/*?:"<>|]', '', title)[:40]
    filename = f"{safe_title}.docx"

    return docx_bytes, filename


def _build_docx(title: str, sections: list, references: list, style: str, search_results: list) -> bytes:
    """用 python-docx 构建 Word 文档"""
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 标题
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 副标题：生成信息
    now = datetime.now(TZ).strftime("%Y年%m月%d日")
    sub = doc.add_paragraph(f"由 Sunday 生成 · {now}", style='Subtitle')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # 正文
    for sec in sections:
        doc.add_heading(sec["heading"], level=2)
        for para_text in sec["body"].split("\n"):
            para_text = para_text.strip()
            if para_text:
                p = doc.add_paragraph(para_text)
                p.style.font.size = Pt(11)

    # 参考文献
    if references:
        doc.add_page_break()
        doc.add_heading("参考文献", level=1)
        for ref in references:
            p = doc.add_paragraph(ref)
            p.style.font.size = Pt(10)

    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph("— 本报告由 Sunday AI 助手自动生成，仅供参考 —")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ============================================================
# 数据图表生成
# ============================================================

def generate_chart(chart_type: str, data: dict, title: str = "") -> bytes:
    """
    用 matplotlib 生成图表 PNG。
    chart_type: bar / line / pie
    data: {"labels": [...], "values": [...]}
    返回 PNG bytes
    """
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt

    plt.rcParams['font.sans-serif'] = ['PingFang SC', 'DejaVu Sans']
    plt.rcParams['axes.unicode_minus'] = False

    labels = data.get("labels", [])
    values = data.get("values", [])

    fig, ax = plt.subplots(figsize=(8, 5))
    colors = ['#ff6b8a', '#ff8fab', '#ffb3c6', '#ffc2d1', '#ffe0e8', '#ffd0da']

    if chart_type == "bar":
        ax.bar(labels, values, color=colors[:len(labels)])
    elif chart_type == "line":
        ax.plot(labels, values, marker='o', color='#ff6b8a', linewidth=2)
    elif chart_type == "pie":
        ax.pie(values, labels=labels, autopct='%1.1f%%', colors=colors[:len(labels)])
    else:
        ax.bar(labels, values, color=colors[:len(labels)])

    if title:
        ax.set_title(title, fontsize=14, fontweight='bold')
    ax.tick_params(labelsize=10)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ============================================================
# 辅助：base64 编码文件
# ============================================================

def encode_attachment(file_bytes: bytes, filename: str) -> dict:
    """将文件编码为 Resend attachment 格式"""
    return {
        "filename": filename,
        "content": base64.b64encode(file_bytes).decode(),
    }
